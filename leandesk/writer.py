from __future__ import annotations

import io
import base64
import os
import re
import tempfile
import uuid
from datetime import datetime
from pathlib import Path
import tkinter as tk
from tkinter import colorchooser, filedialog, font as tkfont, messagebox, ttk

from .core import AppSettings, RecentFiles, RecoveryRecord, RecoveryStore
from .document_formats import LeanDocument, TagRange, read_text_document, write_text_document
from .spellcheck import SpellService, unique_misspelled_words
from .compatibility import WRITER_COMPAT, convert_with_libreoffice, extract_odt_text
from .save_policy import (
    ImportedSourceProtectionError,
    SavePolicyError,
    UnsupportedSaveFormatError,
    imported_source_for,
    mark_save_boundary,
    validate_destination,
    write_atomically,
)
from .ui import COLORS, set_theme_roles


class WriterFrame(ttk.Frame):
    """LeanDesk Writer foundation editor."""

    def __init__(self, master, *, recent: RecentFiles, settings: AppSettings, on_recent_changed=None, on_title_changed=None):
        super().__init__(master)
        self.recent = recent
        self.settings = settings
        self.on_recent_changed = on_recent_changed
        self.on_title_changed = on_title_changed
        self.recovery = RecoveryStore()
        self.recovery_id = str(uuid.uuid4())
        self.current_path: Path | None = None
        self.imported_source_path: Path | None = None
        self.dirty = False
        self.zoom = max(50, min(200, int(settings.default_zoom)))
        self.find_window: tk.Toplevel | None = None
        self.active_ribbon_tab = "Home"
        self.format_painter_tags: tuple[str, ...] = ()
        self.document_metadata: dict[str, object] = {}
        self.writer_objects: dict[str, dict[str, object]] = {}
        self.object_widgets: dict[str, tk.Widget] = {}
        self.ruler_visible = True
        self.line_spacing_var = tk.StringVar(value="1.15")
        self.line_spacing_label = tk.StringVar(value="↕  1.15")
        self._autosave_job = None
        self._spell_job = None
        self.spell_service = SpellService()
        self.live_spell_var = tk.BooleanVar(value=bool(getattr(settings, "live_spellcheck", True)))
        self.last_misspellings: list[tuple[str, str, str]] = []
        self._build_ui()
        self.new_document(confirm=False)
        self._schedule_autosave()

    def _build_ui(self) -> None:
        self.configure(style="WriterRoot.TFrame")
        self._configure_ribbon_styles()

        self.ribbon_shell = tk.Frame(self, bg=COLORS["ribbon"], bd=0, highlightthickness=0)
        self.ribbon_shell.pack(fill="x")

        self.tab_strip = tk.Frame(self.ribbon_shell, bg=COLORS["tab_bg"], height=45)
        self.tab_strip.pack(fill="x")
        self.tab_strip.pack_propagate(False)

        self.file_button = tk.Menubutton(
            self.tab_strip,
            text="File",
            bg=COLORS["tab_bg"],
            fg=COLORS["text"],
            activebackground=COLORS["tab_hover"],
            activeforeground=COLORS["button_active_text"],
            relief="flat",
            bd=0,
            padx=18,
            pady=10,
            font=("Segoe UI Semibold", 10),
            cursor="hand2",
        )
        self.file_button.pack(side="left", padx=(8, 1))
        self.file_menu = tk.Menu(
            self.file_button,
            tearoff=False,
            bg=COLORS["menu_bg"],
            fg=COLORS["menu_text"],
            activebackground=COLORS["menu_active_bg"],
            activeforeground=COLORS["menu_active_text"],
            bd=0,
            relief="flat",
            font=("Segoe UI", 10),
        )
        self.file_menu.add_command(label="New document        Ctrl+N", command=self.new_document)
        self.file_menu.add_command(label="Open...                 Ctrl+O", command=self.open_document)
        self.file_menu.add_separator()
        self.file_menu.add_command(label="Save                    Ctrl+S", command=self.save)
        self.file_menu.add_command(label="Save As...", command=self.save_as)
        self.file_menu.add_command(label="Export PDF...", command=self.export_pdf)
        self.file_menu.add_separator()
        self.file_menu.add_command(label="Print...", command=self.print_document)
        self.file_menu.add_separator()
        self.file_menu.add_command(label="Close document", command=self.new_document)
        self.file_button.configure(menu=self.file_menu)

        self.tab_buttons: dict[str, tk.Button] = {}
        for name in ("Home", "Insert", "Layout", "Review", "View", "Help"):
            button = tk.Button(
                self.tab_strip,
                text=name,
                command=lambda value=name: self._show_ribbon_tab(value),
                bg=COLORS["tab_bg"],
                fg=COLORS["text"],
                activebackground=COLORS["tab_hover"],
                activeforeground=COLORS["button_active_text"],
                relief="flat",
                bd=0,
                padx=16,
                pady=10,
                font=("Segoe UI Semibold", 10),
                cursor="hand2",
            )
            button.pack(side="left", padx=1)
            self.tab_buttons[name] = button

        self.ribbon_body = tk.Frame(
            self.ribbon_shell,
            bg=COLORS["panel"],
            height=150,
            highlightbackground=COLORS["line"],
            highlightthickness=1,
        )
        self.ribbon_body.pack(fill="x")
        self.ribbon_body.pack_propagate(False)

        self._show_ribbon_tab("Home")

        self.ruler = tk.Canvas(self, height=28, bg=COLORS["ruler"], highlightthickness=0)
        self.ruler.pack(fill="x")
        self._draw_ruler()

        editor_shell = tk.Frame(self, bg=COLORS["workspace"])
        editor_shell.pack(fill="both", expand=True)

        self.canvas = tk.Canvas(editor_shell, bg=COLORS["workspace"], highlightthickness=0)
        self.canvas.pack(side="left", fill="both", expand=True)
        yscroll = ttk.Scrollbar(editor_shell, orient="vertical")
        yscroll.pack(side="right", fill="y")

        page = tk.Frame(
            self.canvas,
            bg=COLORS["paper_alt"],
            highlightbackground=COLORS["grid"],
            highlightthickness=1,
        )
        self.page_window = self.canvas.create_window((0, 18), window=page, anchor="n")
        self.text = tk.Text(
            page,
            wrap="word",
            undo=True,
            autoseparators=True,
            maxundo=-1,
            bg=COLORS["paper_alt"],
            fg=COLORS["paper_text"],
            insertbackground=COLORS["paper_text"],
            selectbackground=COLORS["grid"],
            relief="flat",
            borderwidth=0,
            padx=72,
            pady=60,
            font=("Segoe UI", 11),
            spacing1=2,
            spacing2=2,
            spacing3=4,
        )
        set_theme_roles(
            self.text,
            background="paper_alt",
            foreground="paper_text",
            insertbackground="paper_text",
            selectbackground="grid",
        )
        self.text.pack(fill="both", expand=True)
        yscroll.configure(command=self.text.yview)
        self.text.configure(yscrollcommand=yscroll.set, width=95, height=55)
        page.configure(width=920, height=1120)

        self.canvas.bind("<Configure>", self._center_page)
        self.text.bind("<<Modified>>", self._on_modified)
        self.text.bind("<Control-s>", lambda _e: self._shortcut(self.save))
        self.text.bind("<Control-o>", lambda _e: self._shortcut(self.open_document))
        self.text.bind("<Control-n>", lambda _e: self._shortcut(self.new_document))
        self.text.bind("<Control-f>", lambda _e: self._shortcut(self.open_find))
        self.text.bind("<Control-b>", lambda _e: self._shortcut(lambda: self.toggle_tag("fmt_bold")))
        self.text.bind("<Control-i>", lambda _e: self._shortcut(lambda: self.toggle_tag("fmt_italic")))
        self.text.bind("<Control-u>", lambda _e: self._shortcut(lambda: self.toggle_tag("fmt_underline")))
        self.text.bind("<Control-MouseWheel>", self._zoom_wheel)
        self.text.bind("<Button-3>", self.open_spelling_menu)
        self._configure_tags()
        self.text.tag_configure("spell_error", underline=True, foreground="#c9233f")

        self.status_frame = tk.Frame(self, bg=COLORS["status_bg"], height=31)
        self.status_frame.pack(fill="x")
        self.status_frame.pack_propagate(False)
        self.status_var = tk.StringVar(value="Ready")
        self.count_var = tk.StringVar(value="0 words | 0 characters")
        self.path_var = tk.StringVar(value="Unsaved document")
        tk.Label(self.status_frame, textvariable=self.status_var, bg=COLORS["status_bg"], fg=COLORS["text"], font=("Segoe UI", 9)).pack(side="left", padx=10)
        tk.Label(self.status_frame, textvariable=self.count_var, bg=COLORS["status_bg"], fg=COLORS["muted"], font=("Segoe UI", 9)).pack(side="left", padx=20)
        tk.Label(self.status_frame, textvariable=self.path_var, bg=COLORS["status_bg"], fg=COLORS["muted"], font=("Segoe UI", 8)).pack(side="left", padx=12)

        tk.Button(
            self.status_frame,
            text="+",
            command=lambda: self.set_zoom(self.zoom + 10),
            bg=COLORS["status_bg"],
            fg=COLORS["text"],
            activebackground=COLORS["button_hover"],
            activeforeground=COLORS["button_active_text"],
            relief="flat",
            bd=0,
            width=3,
        ).pack(side="right", padx=(0, 6))
        self.zoom_label = tk.StringVar(value=f"{self.zoom}%")
        tk.Label(self.status_frame, textvariable=self.zoom_label, bg=COLORS["status_bg"], fg=COLORS["text"], font=("Segoe UI", 9), width=7).pack(side="right")
        tk.Button(
            self.status_frame,
            text="−",
            command=lambda: self.set_zoom(self.zoom - 10),
            bg=COLORS["status_bg"],
            fg=COLORS["text"],
            activebackground=COLORS["button_hover"],
            activeforeground=COLORS["button_active_text"],
            relief="flat",
            bd=0,
            width=3,
        ).pack(side="right")
        self._update_counts()

    def _configure_ribbon_styles(self) -> None:
        style = ttk.Style(self)
        style.configure(
            "Ribbon.TCombobox",
            fieldbackground=COLORS["panel2"],
            background=COLORS["panel2"],
            foreground=COLORS["text"],
            arrowcolor=COLORS["cobalt"],
            bordercolor=COLORS["line"],
            lightcolor=COLORS["line"],
            darkcolor=COLORS["line"],
            padding=4,
        )
        style.map(
            "Ribbon.TCombobox",
            fieldbackground=[("readonly", COLORS["panel2"])],
            foreground=[("readonly", COLORS["text"])],
        )

    def _draw_ruler(self) -> None:
        self.ruler.delete("all")
        width = max(self.ruler.winfo_width(), 1200)
        self.ruler.create_rectangle(0, 0, width, 28, fill=COLORS["ruler"], outline="")
        start_x = 38
        pixels_per_inch = 70
        for inch in range(0, 19):
            x = start_x + inch * pixels_per_inch
            self.ruler.create_line(x, 8, x, 27, fill=COLORS["ruler_text"], width=1)
            if inch:
                self.ruler.create_text(x + 4, 3, text=str(inch), fill=COLORS["ruler_text"], anchor="nw", font=("Segoe UI", 7))
            for quarter in range(1, 4):
                qx = x + quarter * pixels_per_inch / 4
                length = 9 if quarter == 2 else 6
                self.ruler.create_line(qx, 28 - length, qx, 27, fill=COLORS["ruler_text"])
        self.ruler.create_polygon(110, 4, 104, 12, 116, 12, fill=COLORS["focus"], outline="")
        self.ruler.create_polygon(800, 24, 794, 16, 806, 16, fill=COLORS["focus"], outline="")

    def _tab_button(self, parent, text: str, command, *, italic: bool = False, width: int | None = None) -> tk.Button:
        font = ("Georgia", 12, "italic") if italic else ("Segoe UI", 10)
        button = tk.Button(
            parent,
            text=text,
            command=command,
            bg=COLORS["button_bg"],
            fg=COLORS["button_text"],
            activebackground=COLORS["button_hover"],
            activeforeground=COLORS["button_active_text"],
            relief="flat",
            bd=0,
            padx=9,
            pady=6,
            font=font,
            cursor="hand2",
        )
        if width:
            button.configure(width=width)
        return button

    def _ribbon_group(self, title: str, *, width: int | None = None) -> tuple[tk.Frame, tk.Frame]:
        group = tk.Frame(
            self.ribbon_body,
            bg=COLORS["panel"],
            highlightbackground=COLORS["line"],
            highlightthickness=0,
        )
        group.pack(side="left", fill="y", padx=(7, 0), pady=7)
        if width:
            group.configure(width=width)
            group.pack_propagate(False)
        body = tk.Frame(group, bg=COLORS["panel"])
        body.pack(fill="both", expand=True, padx=4, pady=(2, 0))
        tk.Label(group, text=title, bg=COLORS["panel"], fg=COLORS["muted"], font=("Segoe UI", 8)).pack(side="bottom", pady=(0, 2))
        tk.Frame(self.ribbon_body, bg=COLORS["line"], width=1).pack(side="left", fill="y", pady=14, padx=(7, 0))
        return group, body

    def _show_ribbon_tab(self, name: str) -> None:
        self.active_ribbon_tab = name
        for tab_name, button in self.tab_buttons.items():
            if tab_name == name:
                button.configure(bg=COLORS["button_pressed"], fg=COLORS["button_active_text"])
            else:
                button.configure(bg=COLORS["tab_bg"], fg=COLORS["text"])
        for child in self.ribbon_body.winfo_children():
            child.destroy()
        builders = {
            "Home": self._build_home_ribbon,
            "Insert": self._build_insert_ribbon,
            "Layout": self._build_layout_ribbon,
            "Review": self._build_review_ribbon,
            "View": self._build_view_ribbon,
            "Help": self._build_help_ribbon,
        }
        builders[name]()

    def _build_home_ribbon(self) -> None:
        _, clipboard = self._ribbon_group("Clipboard", width=152)
        paste = self._tab_button(clipboard, "Paste", self.paste_text, width=6)
        paste.grid(row=0, column=0, rowspan=2, sticky="nsew", padx=2, pady=2)
        self._tab_button(clipboard, "Cut", self.cut_text, width=5).grid(row=0, column=1, padx=2, pady=2)
        self._tab_button(clipboard, "Copy", self.copy_text, width=5).grid(row=1, column=1, padx=2, pady=2)
        self._tab_button(clipboard, "Format Painter", self.format_painter, width=13).grid(row=2, column=0, columnspan=2, sticky="ew", padx=2, pady=2)

        _, font_group = self._ribbon_group("Font", width=270)
        self.font_family = tk.StringVar(value=getattr(self, "font_family", tk.StringVar(value="Segoe UI")).get())
        families = sorted(set(tkfont.families()))
        family_box = ttk.Combobox(font_group, textvariable=self.font_family, values=families, state="readonly", width=18, style="Ribbon.TCombobox")
        family_box.grid(row=0, column=0, columnspan=5, sticky="ew", padx=2, pady=2)
        family_box.bind("<<ComboboxSelected>>", lambda _e: self.apply_font_family())
        self.font_size = tk.StringVar(value=getattr(self, "font_size", tk.StringVar(value="11")).get())
        size_box = ttk.Combobox(font_group, textvariable=self.font_size, values=("8","9","10","11","12","14","16","18","20","24","28","32","36","48","60","72"), width=5, style="Ribbon.TCombobox")
        size_box.grid(row=0, column=5, sticky="ew", padx=2, pady=2)
        size_box.bind("<<ComboboxSelected>>", lambda _e: self.apply_font_size())
        size_box.bind("<Return>", lambda _e: self.apply_font_size())
        buttons = [
            ("B", lambda: self.toggle_tag("fmt_bold"), False),
            ("I", lambda: self.toggle_tag("fmt_italic"), True),
            ("U", lambda: self.toggle_tag("fmt_underline"), False),
            ("S̶", lambda: self.toggle_tag("fmt_strike"), False),
            ("x₂", lambda: self.toggle_tag("fmt_subscript"), False),
            ("x²", lambda: self.toggle_tag("fmt_superscript"), False),
            ("A Color", self.choose_color, False),
            ("Highlight", self.choose_highlight, False),
        ]
        for index, (label, command, italic) in enumerate(buttons):
            self._tab_button(font_group, label, command, italic=italic).grid(row=1 + index // 6, column=index % 6, sticky="ew", padx=2, pady=2)

        _, paragraph = self._ribbon_group("Paragraph", width=282)
        paragraph_buttons = [
            ("• List", self.toggle_bullets),
            ("1. List", self.toggle_numbers),
            ("⇤", self.decrease_indent),
            ("⇥", self.increase_indent),
            ("Left", lambda: self.apply_alignment("align_left")),
            ("Center", lambda: self.apply_alignment("align_center")),
            ("Right", lambda: self.apply_alignment("align_right")),
            ("Justify", lambda: self.apply_alignment("align_justify")),
        ]
        for index, (label, command) in enumerate(paragraph_buttons):
            self._tab_button(paragraph, label, command).grid(row=index // 4, column=index % 4, sticky="ew", padx=2, pady=2)
        spacing_button = tk.Menubutton(
            paragraph,
            textvariable=self.line_spacing_label,
            bg=COLORS["button_bg"],
            fg=COLORS["button_text"],
            activebackground=COLORS["button_hover"],
            activeforeground=COLORS["button_active_text"],
            relief="flat",
            bd=0,
            padx=10,
            pady=6,
            font=("Segoe UI", 10),
            cursor="hand2",
        )
        spacing_menu = tk.Menu(spacing_button, tearoff=False, bg=COLORS["menu_bg"], fg=COLORS["menu_text"], activebackground=COLORS["menu_active_bg"], activeforeground=COLORS["menu_active_text"], bd=0)
        for value in ("1.0", "1.15", "1.5", "2.0", "2.5", "3.0"):
            spacing_menu.add_radiobutton(label=value, variable=self.line_spacing_var, value=value, command=lambda value=value: self.apply_line_spacing(value))
        spacing_menu.add_separator()
        spacing_menu.add_command(label="Line Spacing Options...", command=self.line_spacing_options)
        spacing_menu.add_separator()
        spacing_menu.add_command(label="Add Space Before Paragraph", command=lambda: self.apply_paragraph_spacing(before=True))
        spacing_menu.add_command(label="Add Space After Paragraph", command=lambda: self.apply_paragraph_spacing(before=False))
        spacing_menu.add_command(label="Remove Extra Paragraph Space", command=self.remove_paragraph_spacing)
        spacing_button.configure(menu=spacing_menu)
        spacing_button.grid(row=2, column=0, columnspan=4, sticky="ew", padx=2, pady=2)

        _, styles = self._ribbon_group("Styles", width=238)
        style_items = (("Normal", "normal"), ("Heading 1", "heading_1"), ("Heading 2", "heading_2"), ("Heading 3", "heading_3"), ("Title", "heading_1"))
        for index, (label, tag) in enumerate(style_items):
            button = tk.Button(
                styles,
                text=label,
                command=lambda value=tag: self.apply_paragraph_style(value),
                bg=COLORS["button_bg"] if index else COLORS["accent_bg"],
                fg=COLORS["button_text"],
                activebackground=COLORS["button_hover"],
                activeforeground=COLORS["button_active_text"],
                relief="flat",
                bd=0,
                padx=10,
                pady=15,
                font=("Segoe UI", 10 if index else 10, "bold" if index else "normal"),
                cursor="hand2",
            )
            button.grid(row=index // 3, column=index % 3, sticky="nsew", padx=2, pady=2)
            styles.columnconfigure(index % 3, weight=1)

        _, editing = self._ribbon_group("Editing", width=106)
        self._tab_button(editing, "Find", self.open_find, width=12).pack(fill="x", padx=2, pady=2)
        self._tab_button(editing, "Replace", self.open_find, width=12).pack(fill="x", padx=2, pady=2)
        self._tab_button(editing, "Select All", self.select_all, width=12).pack(fill="x", padx=2, pady=2)

    def _build_insert_ribbon(self) -> None:
        _, text_group = self._ribbon_group("Text and symbols", width=390)
        actions = (
            ("Date / Time", self.insert_datetime),
            ("Horizontal Rule", self.insert_rule),
            ("Page Break", self.insert_page_break),
            ("Symbol", self.insert_symbol),
            ("Link", self.insert_link),
            ("Table", self.insert_text_table),
        )
        for index, (label, command) in enumerate(actions):
            self._tab_button(text_group, label, command, width=13).grid(
                row=index // 3, column=index % 3, sticky="ew", padx=3, pady=4
            )
            text_group.columnconfigure(index % 3, weight=1)

        _, media = self._ribbon_group("Media", width=230)
        self._tab_button(media, "Picture", self.insert_image_reference, width=16).pack(
            side="left", padx=4, pady=14
        )
        self._tab_button(media, "File Path", self.insert_file_reference, width=12).pack(
            side="left", padx=4, pady=14
        )

        _, navigation = self._ribbon_group("Navigation", width=240)
        self._tab_button(navigation, "Document Start", lambda: self.text.mark_set("insert", "1.0"), width=14).pack(side="left", padx=3, pady=14)
        self._tab_button(navigation, "Document End", lambda: self.text.mark_set("insert", "end-1c"), width=14).pack(side="left", padx=3, pady=14)
        _, page_parts = self._ribbon_group("Header and footer", width=330)
        for index, (label, command) in enumerate((
            ("Header", self.edit_header), ("Footer", self.edit_footer),
            ("Page Numbers", self.toggle_page_numbers),
        )):
            self._tab_button(page_parts, label, command, width=12).grid(row=0, column=index, padx=3, pady=14)

    def _build_layout_ribbon(self) -> None:
        _, spacing = self._ribbon_group("Spacing", width=350)
        for value in ("1.0", "1.15", "1.5", "2.0", "2.5"):
            self._tab_button(spacing, value, lambda value=value: self.apply_line_spacing(value), width=5).grid(
                row=0, column=(0 if value == "1.0" else ("1.15", "1.5", "2.0", "2.5").index(value) + 1), padx=2, pady=3
            )
        self._tab_button(spacing, "Space Before", lambda: self.apply_paragraph_spacing(before=True), width=12).grid(row=1, column=0, columnspan=2, padx=2, pady=3)
        self._tab_button(spacing, "Space After", lambda: self.apply_paragraph_spacing(before=False), width=12).grid(row=1, column=2, columnspan=2, padx=2, pady=3)
        self._tab_button(spacing, "Clear Extra", self.remove_paragraph_spacing, width=10).grid(row=1, column=4, padx=2, pady=3)

        _, indent = self._ribbon_group("Indent and alignment", width=340)
        for index, (label, command) in enumerate((
            ("Indent −", self.decrease_indent),
            ("Indent +", self.increase_indent),
            ("Left", lambda: self.apply_alignment("align_left")),
            ("Center", lambda: self.apply_alignment("align_center")),
            ("Right", lambda: self.apply_alignment("align_right")),
            ("Justify", lambda: self.apply_alignment("align_justify")),
        )):
            self._tab_button(indent, label, command, width=9).grid(row=index // 3, column=index % 3, padx=3, pady=4)

        _, page = self._ribbon_group("Page view", width=230)
        self._tab_button(page, "Narrow Margins", lambda: self.set_page_margins(42), width=14).pack(side="left", padx=3, pady=14)
        self._tab_button(page, "Normal Margins", lambda: self.set_page_margins(72), width=14).pack(side="left", padx=3, pady=14)
        self._tab_button(page, "Orientation", self.toggle_orientation, width=12).pack(side="left", padx=3, pady=14)

    def _build_review_ribbon(self) -> None:
        _, proofing = self._ribbon_group("Proofing", width=430)
        self._tab_button(proofing, "Check Document", self.check_document_spelling, width=15).grid(row=0, column=0, padx=3, pady=4)
        self._tab_button(proofing, "Next Error", self.next_misspelling, width=12).grid(row=0, column=1, padx=3, pady=4)
        self._tab_button(proofing, "Personal Words", self.manage_personal_dictionary, width=14).grid(row=0, column=2, padx=3, pady=4)
        live = tk.Checkbutton(
            proofing,
            text="Live spell check",
            variable=self.live_spell_var,
            command=self.toggle_live_spellcheck,
            bg=COLORS["panel"],
            fg=COLORS["text"],
            activebackground=COLORS["panel"],
            activeforeground=COLORS["button_active_text"],
            selectcolor=COLORS["panel2"],
            font=("Segoe UI", 9),
        )
        live.grid(row=1, column=0, columnspan=2, sticky="w", padx=5, pady=5)
        tk.Label(
            proofing,
            text=f"Offline dictionary: {self.spell_service.engine_name}",
            bg=COLORS["panel"],
            fg=COLORS["muted"],
            font=("Segoe UI", 8),
        ).grid(row=1, column=2, sticky="w", padx=4)

        _, editing = self._ribbon_group("Editing", width=360)
        for index, (label, command) in enumerate((
            ("Find / Replace", self.open_find),
            ("Word Count", self.show_word_count),
            ("Select All", self.select_all),
            ("Clear Formatting", self.clear_formatting),
        )):
            self._tab_button(editing, label, command, width=14).grid(row=index // 2, column=index % 2, sticky="ew", padx=3, pady=4)

    def _build_view_ribbon(self) -> None:
        _, document = self._ribbon_group("Document view", width=420)
        for index, (label, command) in enumerate((
            ("Toggle Ruler", self.toggle_ruler),
            ("Focus Mode", self.toggle_focus_mode),
            ("Page Width", lambda: self.set_zoom(115)),
            ("75%", lambda: self.set_zoom(75)),
            ("100%", lambda: self.set_zoom(100)),
            ("125%", lambda: self.set_zoom(125)),
        )):
            self._tab_button(document, label, command, width=11).grid(row=index // 3, column=index % 3, padx=3, pady=4)

        _, navigation = self._ribbon_group("Navigation", width=260)
        self._tab_button(navigation, "Top", lambda: self.goto_index("1.0"), width=9).pack(side="left", padx=3, pady=14)
        self._tab_button(navigation, "Bottom", lambda: self.goto_index("end-1c"), width=9).pack(side="left", padx=3, pady=14)
        self._tab_button(navigation, "Find", self.open_find, width=9).pack(side="left", padx=3, pady=14)

    def _build_help_ribbon(self) -> None:
        _, help_group = self._ribbon_group("Writer help", width=440)
        self._tab_button(help_group, "Keyboard Shortcuts", self.show_shortcuts, width=18).pack(side="left", padx=3, pady=14)
        self._tab_button(help_group, "About Writer", self.show_writer_about, width=14).pack(side="left", padx=3, pady=14)
        self._tab_button(help_group, "Data Folder", self.open_writer_data_folder, width=13).pack(side="left", padx=3, pady=14)

        _, formats = self._ribbon_group("Supported formats", width=390)
        tk.Label(
            formats,
            text="Native .ldoc • TXT • Markdown • HTML • RTF • basic DOCX • PDF export",
            bg=COLORS["panel"],
            fg=COLORS["muted"],
            wraplength=350,
            justify="left",
            font=("Segoe UI", 9),
        ).pack(anchor="w", padx=10, pady=20)

    def cut_text(self) -> None:
        try:
            self.text.event_generate("<<Cut>>")
        except tk.TclError:
            pass

    def copy_text(self) -> None:
        try:
            self.text.event_generate("<<Copy>>")
        except tk.TclError:
            pass

    def paste_text(self) -> None:
        try:
            self.text.event_generate("<<Paste>>")
        except tk.TclError:
            pass

    def select_all(self) -> None:
        self.text.tag_add("sel", "1.0", "end-1c")
        self.text.mark_set("insert", "1.0")
        self.text.see("1.0")

    def format_painter(self) -> None:
        if not self.format_painter_tags:
            index = self.text.index("sel.first") if self.text.tag_ranges("sel") else self.text.index("insert")
            self.format_painter_tags = tuple(tag for tag in self.text.tag_names(index) if tag not in {"sel", "search_hit"})
            self.status_var.set("Format Painter armed — select target text and click again")
            return
        if not self.text.tag_ranges("sel"):
            self.status_var.set("Select target text before applying Format Painter")
            return
        start, end = self.text.index("sel.first"), self.text.index("sel.last")
        for tag in self.format_painter_tags:
            self.text.tag_add(tag, start, end)
        self.format_painter_tags = ()
        self.dirty = True
        self.status_var.set("Formatting applied")

    def choose_highlight(self) -> None:
        color = colorchooser.askcolor(initialcolor="#fff08a", parent=self)[1]
        if not color:
            return
        tag = "highlight_" + color.lstrip("#")
        self.text.tag_configure(tag, background=color)
        start, end = self._selection_or_line()
        self.text.tag_add(tag, start, end)
        self.dirty = True

    def increase_indent(self) -> None:
        start, end = self._selection_or_line()
        for tag in tuple(self.text.tag_names(start)):
            if tag.startswith("indent_"):
                self.text.tag_remove(tag, start, end)
        self.text.tag_configure("indent_40", lmargin1=40, lmargin2=40, rmargin=10)
        self.text.tag_add("indent_40", start, end)
        self.dirty = True

    def decrease_indent(self) -> None:
        start, end = self._selection_or_line()
        for tag in self.text.tag_names():
            if tag.startswith("indent_"):
                self.text.tag_remove(tag, start, end)
        self.dirty = True

    def apply_line_spacing(self, value: str) -> None:
        try:
            multiplier = float(value)
        except ValueError:
            multiplier = 1.15
            value = "1.15"
        self.line_spacing_var.set(value)
        self.line_spacing_label.set(f"↕  {value}")
        start, end = self._selection_or_line()
        for tag in self.text.tag_names():
            if tag.startswith("line_spacing_"):
                self.text.tag_remove(tag, start, end)
        spacing_pixels = max(0, round((multiplier - 1.0) * 14))
        tag = "line_spacing_" + value.replace(".", "_")
        self.text.tag_configure(tag, spacing2=spacing_pixels, spacing3=max(2, spacing_pixels // 2))
        self.text.tag_add(tag, start, end)
        self.dirty = True
        self.status_var.set(f"Line spacing set to {value}")

    def line_spacing_options(self) -> None:
        value = simpledialog.askfloat("Line Spacing", "Enter a line-spacing multiplier (0.8 to 4.0):", initialvalue=float(self.line_spacing_var.get()), minvalue=0.8, maxvalue=4.0, parent=self)
        if value is not None:
            self.apply_line_spacing(f"{value:g}")

    def apply_paragraph_spacing(self, *, before: bool) -> None:
        start, end = self._selection_or_line()
        tag = "paragraph_before_10" if before else "paragraph_after_10"
        self.text.tag_configure(tag, spacing1=10 if before else 0, spacing3=10 if not before else 0)
        self.text.tag_add(tag, start, end)
        self.dirty = True
        self.status_var.set("Paragraph spacing added")

    def remove_paragraph_spacing(self) -> None:
        start, end = self._selection_or_line()
        for tag in self.text.tag_names():
            if tag.startswith("paragraph_"):
                self.text.tag_remove(tag, start, end)
        self.dirty = True
        self.status_var.set("Extra paragraph spacing removed")

    def insert_page_break(self) -> None:
        self.text.insert("insert", "\n\f\n")

    def insert_symbol(self) -> None:
        symbol = simpledialog.askstring("Insert Symbol", "Enter a symbol or special character:", parent=self)
        if symbol:
            self.text.insert("insert", symbol)

    def insert_link(self) -> None:
        label = simpledialog.askstring("Insert Link", "Link text:", parent=self)
        if label is None:
            return
        url = simpledialog.askstring("Insert Link", "Address (https://...):", parent=self)
        if url:
            self.text.insert("insert", f"{label or url} <{url}>")

    def insert_text_table(self) -> None:
        rows = simpledialog.askinteger("Insert Table", "Rows:", initialvalue=3, minvalue=1, maxvalue=20, parent=self)
        if rows is None:
            return
        cols = simpledialog.askinteger("Insert Table", "Columns:", initialvalue=3, minvalue=1, maxvalue=12, parent=self)
        if cols is None:
            return
        self.insert_table(rows, cols)

    def insert_table(self, rows: int, cols: int, data: list[list[str]] | None = None, *, index: str = "insert") -> str:
        rows = max(1, min(50, int(rows)))
        cols = max(1, min(20, int(cols)))
        values = data or [["" for _ in range(cols)] for _ in range(rows)]
        object_id = uuid.uuid4().hex
        record: dict[str, object] = {"id": object_id, "kind": "table", "rows": rows, "cols": cols, "data": values}
        self.writer_objects[object_id] = record
        self._place_writer_object(record, index)
        self.dirty = True
        self.status_var.set(f"Inserted {rows} x {cols} table")
        return object_id

    def insert_image_reference(self) -> None:
        value = filedialog.askopenfilename(
            parent=self,
            filetypes=(("Images", "*.png *.jpg *.jpeg *.gif *.bmp *.webp"),),
        )
        if value:
            try:
                self.insert_image(Path(value))
            except Exception as exc:
                messagebox.showerror("Insert Picture", f"Could not insert picture.\n\n{exc}", parent=self)

    def insert_image(self, path: Path, *, index: str = "insert") -> str:
        data = Path(path).read_bytes()
        if len(data) > 16 * 1024 * 1024:
            raise ValueError("Picture exceeds the 16 MB document limit")
        object_id = uuid.uuid4().hex
        record: dict[str, object] = {
            "id": object_id,
            "kind": "image",
            "name": Path(path).name,
            "data_base64": base64.b64encode(data).decode("ascii"),
        }
        self.writer_objects[object_id] = record
        self._place_writer_object(record, index)
        self.dirty = True
        self.status_var.set(f"Inserted {Path(path).name}")
        return object_id

    def _place_writer_object(self, record: dict[str, object], index: str) -> None:
        object_id = str(record["id"])
        mark = f"writer_object_{object_id}"
        self.text.mark_set(mark, index)
        self.text.mark_gravity(mark, "left")
        if record.get("kind") == "table":
            frame = tk.Frame(self.text, bg=COLORS["paper_alt"], highlightbackground=COLORS["grid"], highlightthickness=1)
            variables: list[list[tk.StringVar]] = []
            source = record.get("data", [])
            for row in range(int(record.get("rows", 1))):
                variable_row: list[tk.StringVar] = []
                for col in range(int(record.get("cols", 1))):
                    initial = ""
                    if isinstance(source, list) and row < len(source) and isinstance(source[row], list) and col < len(source[row]):
                        initial = str(source[row][col])
                    variable = tk.StringVar(value=initial)
                    ttk.Entry(frame, textvariable=variable, width=14).grid(row=row, column=col, sticky="nsew", padx=1, pady=1)
                    frame.columnconfigure(col, weight=1)
                    variable_row.append(variable)
                variables.append(variable_row)
            record["_variables"] = variables
            widget: tk.Widget = frame
        else:
            raw = base64.b64decode(str(record.get("data_base64", "")), validate=True)
            from PIL import Image, ImageTk
            image = Image.open(io.BytesIO(raw))
            image.thumbnail((640, 420))
            photo = ImageTk.PhotoImage(image)
            label = tk.Label(self.text, image=photo, bg=COLORS["paper_alt"], bd=1, relief="solid")
            label.image = photo
            widget = label
        self.text.window_create(index, window=widget, padx=8, pady=8)
        self.object_widgets[object_id] = widget

    def set_header(self, value: str) -> None:
        self.document_metadata["header"] = value
        self.dirty = True

    def set_footer(self, value: str) -> None:
        self.document_metadata["footer"] = value
        self.dirty = True

    def edit_header(self) -> None:
        value = simpledialog.askstring("Header", "Header text:", initialvalue=str(self.document_metadata.get("header", "")), parent=self)
        if value is not None:
            self.set_header(value)

    def edit_footer(self) -> None:
        value = simpledialog.askstring("Footer", "Footer text:", initialvalue=str(self.document_metadata.get("footer", "")), parent=self)
        if value is not None:
            self.set_footer(value)

    def toggle_page_numbers(self) -> None:
        enabled = not bool(self.document_metadata.get("page_numbers", False))
        self.document_metadata["page_numbers"] = enabled
        self.dirty = True
        self.status_var.set("Page numbers enabled" if enabled else "Page numbers disabled")

    def toggle_orientation(self) -> None:
        current = str(self.document_metadata.get("orientation", "portrait"))
        self.document_metadata["orientation"] = "landscape" if current == "portrait" else "portrait"
        self.dirty = True
        self.status_var.set(f"Page orientation: {self.document_metadata['orientation']}")

    def insert_file_reference(self) -> None:
        value = filedialog.askopenfilename(parent=self)
        if value:
            self.text.insert("insert", f"[File: {value}]")

    def set_page_margins(self, pixels: int) -> None:
        self.text.configure(padx=max(20, int(pixels)))
        self.document_metadata["margin_pixels"] = max(20, int(pixels))
        self.status_var.set(f"Page margins set to {pixels}px")

    def clear_formatting(self) -> None:
        start, end = self._selection_or_line()
        prefixes = (
            "fmt_", "align_", "heading_", "font_", "size_", "color_",
            "highlight_", "line_spacing_", "paragraph_", "indent_",
        )
        for tag in self.text.tag_names():
            if tag.startswith(prefixes):
                self.text.tag_remove(tag, start, end)
        self.dirty = True
        self._update_title_indicator()

    def toggle_focus_mode(self) -> None:
        if self.ribbon_shell.winfo_ismapped():
            self.ribbon_shell.pack_forget()
            if self.ruler_visible:
                self.ruler.pack_forget()
            self.status_var.set("Focus mode — press View > Focus Mode after restoring the ribbon with F11")
            self.winfo_toplevel().bind("<F11>", lambda _e: self.toggle_focus_mode())
        else:
            self.ribbon_shell.pack(fill="x", before=self.canvas.master)
            if self.ruler_visible:
                self.ruler.pack(fill="x", before=self.canvas.master)
            self.status_var.set("Focus mode closed")

    def goto_index(self, index: str) -> None:
        self.text.mark_set("insert", index)
        self.text.see(index)
        self.text.focus_set()

    def open_writer_data_folder(self) -> None:
        from .core import DATA_ROOT
        DATA_ROOT.mkdir(parents=True, exist_ok=True)
        if os.name == "nt":
            os.startfile(DATA_ROOT)
        else:
            import webbrowser
            webbrowser.open(DATA_ROOT.as_uri())

    def toggle_live_spellcheck(self) -> None:
        self.settings.live_spellcheck = self.live_spell_var.get()
        self.settings.save()
        if self.live_spell_var.get():
            self.schedule_spellcheck(immediate=True)
        else:
            self.text.tag_remove("spell_error", "1.0", "end")
            self.last_misspellings = []
            self.status_var.set("Live spell check off")

    def schedule_spellcheck(self, immediate: bool = False) -> None:
        if self._spell_job:
            try:
                self.after_cancel(self._spell_job)
            except Exception:
                pass
        if not self.live_spell_var.get():
            return
        self._spell_job = self.after(25 if immediate else 650, self.run_spellcheck)

    def run_spellcheck(self) -> int:
        self._spell_job = None
        content = self.text.get("1.0", "end-1c")
        self.text.tag_remove("spell_error", "1.0", "end")
        rows = self.spell_service.misspellings(content)
        self.last_misspellings = []
        for row in rows:
            start = f"1.0+{row.start}c"
            end = f"1.0+{row.end}c"
            self.text.tag_add("spell_error", start, end)
            self.last_misspellings.append((row.word, start, end))
        if rows:
            self.status_var.set(f"Spell check: {len(rows)} possible error{'s' if len(rows) != 1 else ''}")
        else:
            self.status_var.set("Spell check complete — no errors found")
        return len(rows)

    def check_document_spelling(self) -> None:
        count = self.run_spellcheck()
        unique = unique_misspelled_words(self.text.get("1.0", "end-1c"), self.spell_service)
        if not count:
            messagebox.showinfo(
                "Spell Check",
                f"No spelling errors were found.\n\nDictionary: {self.spell_service.engine_name}\nPersonal words: {len(self.spell_service.personal_words)}",
                parent=self,
            )
            return
        preview = "\n".join(f"• {word}" for word in unique[:20])
        extra = f"\n…and {len(unique) - 20} more unique words" if len(unique) > 20 else ""
        messagebox.showinfo(
            "Spell Check",
            f"Found {count} possible spelling errors across {len(unique)} unique words.\n\n{preview}{extra}\n\nRight-click an underlined word for suggestions, or choose Next Error.",
            parent=self,
        )

    def next_misspelling(self) -> None:
        if not self.last_misspellings:
            self.run_spellcheck()
        if not self.last_misspellings:
            return
        current = self.text.index("insert")
        target = None
        for row in self.last_misspellings:
            if self.text.compare(row[1], ">", current):
                target = row
                break
        if target is None:
            target = self.last_misspellings[0]
        word, start, end = target
        self.text.tag_remove("sel", "1.0", "end")
        self.text.tag_add("sel", start, end)
        self.text.mark_set("insert", end)
        self.text.see(start)
        suggestions = self.spell_service.suggestions(word, 5)
        self.status_var.set(f"{word}: " + (", ".join(suggestions) if suggestions else "no suggestions"))

    def open_spelling_menu(self, event) -> str:
        index = self.text.index(f"@{event.x},{event.y}")
        start = self.text.index(f"{index} wordstart")
        end = self.text.index(f"{index} wordend")
        word = self.text.get(start, end).strip()
        menu = tk.Menu(
            self.text,
            tearoff=False,
            bg=COLORS["menu_bg"],
            fg=COLORS["menu_text"],
            activebackground=COLORS["menu_active_bg"],
            activeforeground=COLORS["menu_active_text"],
        )
        if word and not self.spell_service.is_correct(word):
            suggestions = self.spell_service.suggestions(word, 8)
            if suggestions:
                for suggestion in suggestions:
                    menu.add_command(
                        label=suggestion,
                        command=lambda value=suggestion, a=start, b=end: self.replace_spelling(a, b, value),
                    )
            else:
                menu.add_command(label="No suggestions", state="disabled")
            menu.add_separator()
            menu.add_command(label=f'Add "{word}" to dictionary', command=lambda: self.add_word_to_dictionary(word))
            menu.add_command(label="Ignore for this document", command=lambda a=start, b=end: self.text.tag_remove("spell_error", a, b))
        else:
            menu.add_command(label="Cut", command=self.cut_text)
            menu.add_command(label="Copy", command=self.copy_text)
            menu.add_command(label="Paste", command=self.paste_text)
        menu.tk_popup(event.x_root, event.y_root)
        return "break"

    def replace_spelling(self, start: str, end: str, suggestion: str) -> None:
        original = self.text.get(start, end)
        if original[:1].isupper():
            suggestion = suggestion[:1].upper() + suggestion[1:]
        self.text.delete(start, end)
        self.text.insert(start, suggestion)
        self.dirty = True
        self.schedule_spellcheck(immediate=True)

    def add_word_to_dictionary(self, word: str) -> None:
        self.spell_service.add_personal(word)
        self.schedule_spellcheck(immediate=True)
        self.status_var.set(f'Added "{word}" to personal dictionary')

    def manage_personal_dictionary(self) -> None:
        window = tk.Toplevel(self)
        window.title("Personal Dictionary")
        window.geometry("430x430")
        window.configure(bg=COLORS["panel"])
        tk.Label(window, text="PERSONAL DICTIONARY", bg=COLORS["panel"], fg=COLORS["text"], font=("Segoe UI Semibold", 14)).pack(anchor="w", padx=14, pady=(14, 6))
        tk.Label(window, text="Words added here are stored locally and used by live spell checking.", bg=COLORS["panel"], fg=COLORS["muted"], wraplength=390, justify="left").pack(anchor="w", padx=14, pady=(0, 8))
        listbox = tk.Listbox(window, bg=COLORS["field"], fg=COLORS["field_text"], selectbackground=COLORS["selection"], relief="flat")
        listbox.pack(fill="both", expand=True, padx=14, pady=6)
        def refresh():
            listbox.delete(0, "end")
            for value in sorted(self.spell_service.personal_words):
                listbox.insert("end", value)
        def add():
            value = simpledialog.askstring("Personal Dictionary", "Word:", parent=window)
            if value:
                self.spell_service.add_personal(value)
                refresh()
                self.schedule_spellcheck(immediate=True)
        def remove():
            selection = listbox.curselection()
            if selection:
                self.spell_service.remove_personal(listbox.get(selection[0]))
                refresh()
                self.schedule_spellcheck(immediate=True)
        buttons = tk.Frame(window, bg=COLORS["panel"])
        buttons.pack(fill="x", padx=14, pady=(4, 14))
        self._tab_button(buttons, "Add Word", add, width=12).pack(side="left", padx=3)
        self._tab_button(buttons, "Remove", remove, width=12).pack(side="left", padx=3)
        self._tab_button(buttons, "Close", window.destroy, width=12).pack(side="right", padx=3)
        refresh()

    def show_word_count(self) -> None:
        content = self.text.get("1.0", "end-1c")
        words = len(re.findall(r"\b\w+[\w'-]*\b", content, flags=re.UNICODE))
        lines = int(self.text.index("end-1c").split(".")[0])
        messagebox.showinfo("Word Count", f"Words: {words:,}\nCharacters: {len(content):,}\nParagraph lines: {lines:,}", parent=self)

    def toggle_ruler(self) -> None:
        if self.ruler_visible:
            self.ruler.pack_forget()
        else:
            self.ruler.pack(fill="x", before=self.canvas.master)
        self.ruler_visible = not self.ruler_visible

    def set_zoom(self, value: int) -> None:
        self.zoom = max(50, min(200, int(value)))
        size = max(6, round(11 * self.zoom / 100))
        self.text.configure(font=(self.font_family.get(), size))
        self.zoom_label.set(f"{self.zoom}%")
        self.status_var.set(f"Zoom {self.zoom}%")

    def show_shortcuts(self) -> None:
        messagebox.showinfo("Writer Shortcuts", "Ctrl+N  New\nCtrl+O  Open\nCtrl+S  Save\nCtrl+F  Find\nCtrl+B  Bold\nCtrl+I  Italic\nCtrl+U  Underline\nCtrl+mouse wheel  Zoom", parent=self)

    def show_writer_about(self) -> None:
        messagebox.showinfo("About LeanDesk Writer", "LeanDesk Writer 0.9.0\n\nClean grouped ribbon, offline spell checking with personal dictionary, paragraph spacing, recovery, DOCX workflows, and PDF export.", parent=self)

    def export_pdf(self) -> bool:
        value = filedialog.asksaveasfilename(parent=self, defaultextension=".pdf", filetypes=(("PDF", "*.pdf"),))
        if not value:
            return False
        return self._write_to(Path(value))

    @staticmethod
    def _shortcut(command):
        command()
        return "break"

    def _configure_tags(self) -> None:
        self.text.tag_configure("fmt_bold", font=("Segoe UI", 11, "bold"))
        self.text.tag_configure("fmt_italic", font=("Segoe UI", 11, "italic"))
        self.text.tag_configure("fmt_underline", underline=True)
        self.text.tag_configure("fmt_strike", overstrike=True)
        self.text.tag_configure("fmt_subscript", offset=-4, font=("Segoe UI", 8))
        self.text.tag_configure("fmt_superscript", offset=5, font=("Segoe UI", 8))
        self.text.tag_configure("align_left", justify="left")
        self.text.tag_configure("align_center", justify="center")
        self.text.tag_configure("align_right", justify="right")
        self.text.tag_configure("align_justify", justify="left")
        self.text.tag_configure("heading_1", font=("Segoe UI", 24, "bold"), spacing1=10, spacing3=8)
        self.text.tag_configure("heading_2", font=("Segoe UI", 18, "bold"), spacing1=8, spacing3=6)
        self.text.tag_configure("heading_3", font=("Segoe UI", 14, "bold"), spacing1=6, spacing3=4)

    def _center_page(self, event) -> None:
        self.canvas.coords(self.page_window, max(event.width // 2, 450), 18)

    def _zoom_wheel(self, event) -> str:
        self.set_zoom(self.zoom + (10 if event.delta > 0 else -10))
        return "break"

    def _on_modified(self, _event=None) -> None:
        if self.text.edit_modified():
            self.dirty = True
            self.text.edit_modified(False)
            self._update_counts()
            self._update_title_indicator()
            self.schedule_spellcheck()

    def _update_counts(self) -> None:
        content = self.text.get("1.0", "end-1c")
        words = len(re.findall(r"\b\w+[\w'-]*\b", content, flags=re.UNICODE))
        self.count_var.set(f"{words:,} words | {len(content):,} characters")

    def _update_title_indicator(self) -> None:
        title = self.current_path.name if self.current_path else "Untitled Document"
        self.status_var.set(("* " if self.dirty else "") + title)
        if self.on_title_changed:
            self.on_title_changed(title, self.dirty)

    def _selection_or_line(self) -> tuple[str, str]:
        try:
            return self.text.index("sel.first"), self.text.index("sel.last")
        except tk.TclError:
            return self.text.index("insert linestart"), self.text.index("insert lineend+1c")

    def toggle_tag(self, tag: str) -> None:
        start, end = self._selection_or_line()
        if tag in self.text.tag_names(start):
            self.text.tag_remove(tag, start, end)
        else:
            self.text.tag_add(tag, start, end)
        self.dirty = True
        self._update_title_indicator()

    def apply_font_family(self) -> None:
        family = self.font_family.get().strip() or "Segoe UI"
        tag = "font_" + re.sub(r"[^a-zA-Z0-9]+", "_", family)
        self.text.tag_configure(tag, font=(family, int(self.font_size.get() or 11)))
        start, end = self._selection_or_line()
        self.text.tag_add(tag, start, end)
        self.dirty = True

    def apply_font_size(self) -> None:
        try:
            size = max(6, min(144, int(self.font_size.get())))
        except ValueError:
            size = 11
            self.font_size.set("11")
        tag = f"size_{size}"
        self.text.tag_configure(tag, font=(self.font_family.get(), size))
        start, end = self._selection_or_line()
        self.text.tag_add(tag, start, end)
        self.dirty = True

    def choose_color(self) -> None:
        color = colorchooser.askcolor(parent=self)[1]
        if not color:
            return
        tag = "color_" + color.lstrip("#")
        self.text.tag_configure(tag, foreground=color)
        start, end = self._selection_or_line()
        self.text.tag_add(tag, start, end)
        self.dirty = True

    def apply_alignment(self, tag: str) -> None:
        start, end = self._selection_or_line()
        for other in ("align_left", "align_center", "align_right", "align_justify"):
            self.text.tag_remove(other, start, end)
        self.text.tag_add(tag, start, end)
        self.dirty = True

    def apply_paragraph_style(self, tag: str) -> None:
        start = self.text.index("sel.first linestart") if self.text.tag_ranges("sel") else self.text.index("insert linestart")
        end = self.text.index("sel.last lineend+1c") if self.text.tag_ranges("sel") else self.text.index("insert lineend+1c")
        for heading in ("heading_1", "heading_2", "heading_3"):
            self.text.tag_remove(heading, start, end)
        if tag != "normal":
            self.text.tag_add(tag, start, end)
        self.dirty = True

    def _selected_lines(self) -> tuple[str, str]:
        if self.text.tag_ranges("sel"):
            return self.text.index("sel.first linestart"), self.text.index("sel.last lineend")
        return self.text.index("insert linestart"), self.text.index("insert lineend")

    def toggle_bullets(self) -> None:
        start, end = self._selected_lines()
        line = start
        while self.text.compare(line, "<=", end):
            current = self.text.get(line, f"{line} lineend")
            if current.startswith("• "):
                self.text.delete(line, f"{line}+2c")
            else:
                self.text.insert(line, "• ")
            next_line = self.text.index(f"{line}+1line")
            if self.text.compare(next_line, ">", end):
                break
            line = next_line
        self.dirty = True

    def toggle_numbers(self) -> None:
        start, end = self._selected_lines()
        line = start
        number = 1
        while self.text.compare(line, "<=", end):
            current = self.text.get(line, f"{line} lineend")
            cleaned = re.sub(r"^\d+\.\s+", "", current)
            if cleaned != current:
                self.text.delete(line, f"{line} lineend")
                self.text.insert(line, cleaned)
            else:
                self.text.insert(line, f"{number}. ")
            number += 1
            next_line = self.text.index(f"{line}+1line")
            if self.text.compare(next_line, ">", end):
                break
            line = next_line
        self.dirty = True

    def insert_datetime(self) -> None:
        self.text.insert("insert", datetime.now().strftime("%B %d, %Y %I:%M %p"))

    def insert_rule(self) -> None:
        self.text.insert("insert", "\n" + "─" * 62 + "\n")

    def serialize(self) -> LeanDocument:
        tags: list[TagRange] = []
        prefixes = ("fmt_", "align_", "heading_", "font_", "size_", "color_", "highlight_", "line_spacing_", "paragraph_", "indent_")
        for tag in self.text.tag_names():
            if not tag.startswith(prefixes):
                continue
            ranges = self.text.tag_ranges(tag)
            for index in range(0, len(ranges), 2):
                tags.append(TagRange(tag, str(ranges[index]), str(ranges[index + 1])))
        objects: list[dict[str, object]] = []
        for object_id, source in self.writer_objects.items():
            row = {name: value for name, value in source.items() if not name.startswith("_")}
            variables = source.get("_variables")
            if isinstance(variables, list):
                row["data"] = [[variable.get() for variable in variable_row] for variable_row in variables]
            try:
                row["index"] = self.text.index(f"writer_object_{object_id}")
            except tk.TclError:
                row["index"] = "end-1c"
            objects.append(row)
        metadata = dict(self.document_metadata)
        metadata.update({"saved_at": datetime.now().isoformat(timespec="seconds"), "zoom": self.zoom, "objects": objects})
        return LeanDocument(
            title=self.current_path.stem if self.current_path else "Untitled Document",
            text=self.text.get("1.0", "end-1c"),
            tags=tags,
            metadata=metadata,
        )

    def load_document(self, document: LeanDocument, path: Path | None = None) -> None:
        self.text.delete("1.0", "end")
        self.text.insert("1.0", document.text)
        self.document_metadata = dict(document.metadata)
        self.writer_objects = {}
        self.object_widgets = {}
        for row in document.tags:
            if row.tag.startswith("size_"):
                try:
                    self.text.tag_configure(row.tag, font=("Segoe UI", int(row.tag.split("_", 1)[1])))
                except ValueError:
                    pass
            elif row.tag.startswith("color_"):
                self.text.tag_configure(row.tag, foreground="#" + row.tag.split("_", 1)[1])
            elif row.tag.startswith("font_"):
                self.text.tag_configure(row.tag, font=(row.tag.split("_", 1)[1].replace("_", " "), 11))
            elif row.tag.startswith("highlight_"):
                self.text.tag_configure(row.tag, background="#" + row.tag.split("_", 1)[1])
            elif row.tag.startswith("line_spacing_"):
                try:
                    multiplier = float(row.tag.replace("line_spacing_", "").replace("_", "."))
                    spacing_pixels = max(0, round((multiplier - 1.0) * 14))
                    self.text.tag_configure(row.tag, spacing2=spacing_pixels, spacing3=max(2, spacing_pixels // 2))
                except ValueError:
                    pass
            elif row.tag == "paragraph_before_10":
                self.text.tag_configure(row.tag, spacing1=10)
            elif row.tag == "paragraph_after_10":
                self.text.tag_configure(row.tag, spacing3=10)
            elif row.tag.startswith("indent_"):
                self.text.tag_configure(row.tag, lmargin1=40, lmargin2=40, rmargin=10)
            try:
                self.text.tag_add(row.tag, row.start, row.end)
            except tk.TclError:
                pass
        objects = document.metadata.get("objects", [])
        if isinstance(objects, list):
            for raw in objects:
                if not isinstance(raw, dict) or raw.get("kind") not in {"table", "image"}:
                    continue
                record = dict(raw)
                object_id = str(record.get("id") or uuid.uuid4().hex)
                record["id"] = object_id
                self.writer_objects[object_id] = record
                try:
                    self._place_writer_object(record, str(record.get("index", "end-1c")))
                except (ValueError, tk.TclError, OSError):
                    self.writer_objects.pop(object_id, None)
        margin = document.metadata.get("margin_pixels")
        if isinstance(margin, int):
            self.text.configure(padx=max(20, min(180, margin)))
        self.current_path = path
        self.imported_source_path = imported_source_for("Writer", path)
        self.dirty = False
        self.text.edit_modified(False)
        self.path_var.set(str(path) if path else "Unsaved document")
        self._update_counts()
        self._update_title_indicator()
        self.schedule_spellcheck(immediate=True)

    def new_document(self, confirm: bool = True) -> bool:
        if confirm and not self.confirm_discard_or_save():
            return False
        self.recovery.delete(self.recovery_id)
        self.recovery_id = str(uuid.uuid4())
        self.load_document(LeanDocument())
        return True

    def confirm_discard_or_save(self) -> bool:
        if not self.dirty:
            return True
        answer = messagebox.askyesnocancel("LeanDesk Writer", "Save changes before continuing?", parent=self)
        if answer is None:
            return False
        if answer:
            return self.save()
        return True

    def open_document(self, path: Path | None = None) -> bool:
        if not self.confirm_discard_or_save():
            return False
        if path is None:
            value = filedialog.askopenfilename(
                parent=self,
                filetypes=(
                    ("LeanDesk documents", "*.ldoc"),
                    ("Word/OpenDocument", "*.docx *.doc *.docm *.dot *.dotx *.odt *.ott"),
                    ("Legacy/other documents", "*.rtf *.wps *.wpd *.abw *.sxw *.lwp *.cwk *.pages"),
                    ("Text and Markdown", "*.txt *.md"),
                    ("HTML", "*.html *.htm"),
                    ("All files", "*.*"),
                ),
            )
            if not value:
                return False
            path = Path(value)
        try:
            suffix = path.suffix.lower()
            compatibility_note = None
            if suffix == ".docx":
                document = self._load_docx(path)
            elif suffix == ".odt":
                try:
                    converted = convert_with_libreoffice(path, "Writer")
                    document = self._load_docx(converted.as_file(), title=path.stem)
                    compatibility_note = converted.note
                except RuntimeError:
                    document = LeanDocument(title=path.stem, text=extract_odt_text(path))
                    compatibility_note = "ODT opened using reduced-fidelity built-in fallback; original was not modified."
            elif suffix in WRITER_COMPAT:
                converted = convert_with_libreoffice(path, "Writer")
                document = self._load_docx(converted.as_file())
                compatibility_note = converted.note
            else:
                document = read_text_document(path)
        except Exception as exc:
            messagebox.showerror("LeanDesk Writer", f"Could not open the document.\n\n{exc}", parent=self)
            return False
        self.load_document(document, path)
        if compatibility_note:
            self.status_var.set(compatibility_note)
        self.recent.add(path, "Writer")
        if self.on_recent_changed:
            self.on_recent_changed()
        self.recovery.delete(self.recovery_id)
        self.recovery_id = str(uuid.uuid4())
        return True

    def _protected_import_source(self) -> Path | None:
        explicit = getattr(self, "imported_source_path", None)
        if explicit is not None:
            return Path(explicit)
        return imported_source_for("Writer", getattr(self, "current_path", None))

    @mark_save_boundary
    def save(self) -> bool:
        if self._protected_import_source() is not None:
            try:
                proceed = messagebox.askyesno(
                    "Original File Protected",
                    "This document was imported from another format. LeanDesk will not overwrite the original because unsupported features could be lost.\n\nSave a new copy instead?",
                    parent=self,
                )
            except Exception:
                return False
            return self.save_as() if proceed else False
        return self.save_as() if self.current_path is None else self._write_to(self.current_path)

    @mark_save_boundary
    def save_as(self) -> bool:
        value = filedialog.asksaveasfilename(
            parent=self,
            defaultextension=".ldoc",
            filetypes=(
                ("LeanDesk document", "*.ldoc"),
                ("Word document", "*.docx"),
                ("Text", "*.txt"),
                ("Markdown", "*.md"),
                ("HTML", "*.html"),
                ("Rich Text", "*.rtf"),
                ("PDF", "*.pdf"),
            ),
        )
        return bool(value) and self._write_to(Path(value))

    @mark_save_boundary
    def _write_to(self, path: Path) -> bool:
        document = self.serialize()
        try:
            destination = validate_destination(
                "Writer",
                path,
                imported_source=self._protected_import_source(),
                allow_export_only=True,
            )

            def produce(temporary: Path) -> None:
                suffix = destination.suffix.lower()
                if suffix == ".docx":
                    self._save_docx(document, temporary)
                elif suffix == ".pdf":
                    self._save_pdf(document, temporary)
                else:
                    write_text_document(document, temporary)

            write_atomically(destination, produce)
        except (ImportedSourceProtectionError, UnsupportedSaveFormatError, SavePolicyError) as exc:
            messagebox.showwarning("LeanDesk Writer", str(exc), parent=self)
            return False
        except Exception as exc:
            messagebox.showerror("LeanDesk Writer", f"Could not save the document.\n\n{exc}", parent=self)
            return False
        if destination.suffix.lower() == ".pdf":
            self.status_var.set(f"Exported {destination.name}; document source unchanged")
            return True
        self.current_path = destination
        self.imported_source_path = None
        self.dirty = False
        self.path_var.set(str(destination))
        self._update_title_indicator()
        self.recent.add(destination, "Writer")
        if self.on_recent_changed:
            self.on_recent_changed()
        self.recovery.delete(self.recovery_id)
        return True

    @staticmethod
    def _load_docx(path, *, title: str | None = None) -> LeanDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX support requires python-docx. Run the Windows builder or install python-docx.") from exc
        from .ooxml_preflight import prepare_ooxml
        prepared = prepare_ooxml(path, "docx")
        doc = Document(prepared.open())
        lines, tags = [], []
        for line_no, paragraph in enumerate(doc.paragraphs, start=1):
            lines.append(paragraph.text)
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                level = paragraph.style.name.split()[-1]
                if level in {"1", "2", "3"}:
                    tags.append(TagRange(f"heading_{level}", f"{line_no}.0", f"{line_no}.end"))
        if title is None:
            title = Path(path).stem if isinstance(path, (str, Path)) else "Imported Document"
        if not isinstance(title, str) or not title:
            title = "Imported Document"
        objects: list[dict[str, object]] = []
        for table in doc.tables:
            data = [[cell.text for cell in row.cells] for row in table.rows]
            objects.append({"id": uuid.uuid4().hex, "kind": "table", "rows": len(data), "cols": max((len(row) for row in data), default=1), "data": data, "index": "end-1c"})
        seen_images: set[str] = set()
        for relationship in doc.part.rels.values():
            if "image" not in relationship.reltype:
                continue
            blob = relationship.target_part.blob
            digest = str(hash(blob))
            if digest in seen_images or len(blob) > 16 * 1024 * 1024:
                continue
            seen_images.add(digest)
            objects.append({"id": uuid.uuid4().hex, "kind": "image", "name": Path(relationship.target_ref).name, "data_base64": base64.b64encode(blob).decode("ascii"), "index": "end-1c"})
        section = doc.sections[0]
        metadata: dict[str, object] = {
            "header": "\n".join(paragraph.text for paragraph in section.header.paragraphs).strip(),
            "footer": "\n".join(paragraph.text for paragraph in section.footer.paragraphs).strip(),
            "orientation": "landscape" if section.page_width > section.page_height else "portrait",
            "objects": objects,
        }
        return LeanDocument(title=title, text="\n".join(lines), tags=tags, metadata=metadata)

    @staticmethod
    def _save_docx(document: LeanDocument, path: Path) -> None:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX support requires python-docx. Run the Windows builder or install python-docx.") from exc
        from docx.enum.section import WD_ORIENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches
        doc = Document()
        section = doc.sections[0]
        if document.metadata.get("orientation") == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        margin_pixels = document.metadata.get("margin_pixels", 72)
        if isinstance(margin_pixels, int):
            margin = Inches(max(0.25, min(2.0, margin_pixels / 96)))
            section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = margin
        header = str(document.metadata.get("header", ""))
        footer = str(document.metadata.get("footer", ""))
        if header:
            section.header.paragraphs[0].text = header
        if footer:
            section.footer.paragraphs[0].text = footer
        if document.metadata.get("page_numbers"):
            paragraph = section.footer.paragraphs[0]
            if footer:
                paragraph.add_run("  ")
            run = paragraph.add_run()
            field = OxmlElement("w:fldSimple")
            field.set(qn("w:instr"), "PAGE")
            run._r.append(field)
        for line in document.text.splitlines() or [""]:
            doc.add_paragraph(line)
        objects = document.metadata.get("objects", [])
        if isinstance(objects, list):
            for item in objects:
                if not isinstance(item, dict):
                    continue
                if item.get("kind") == "table":
                    rows, cols = int(item.get("rows", 1)), int(item.get("cols", 1))
                    table = doc.add_table(rows=max(1, rows), cols=max(1, cols))
                    table.style = "Table Grid"
                    data = item.get("data", [])
                    for row in range(rows):
                        for col in range(cols):
                            if isinstance(data, list) and row < len(data) and isinstance(data[row], list) and col < len(data[row]):
                                table.cell(row, col).text = str(data[row][col])
                elif item.get("kind") == "image":
                    try:
                        payload = base64.b64decode(str(item.get("data_base64", "")), validate=True)
                        doc.add_picture(io.BytesIO(payload), width=Inches(6.0))
                    except Exception:
                        continue
        doc.save(str(path))

    @staticmethod
    def _save_pdf(document: LeanDocument, path: Path) -> None:
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
            from xml.sax.saxutils import escape
        except ImportError as exc:
            raise RuntimeError("PDF export requires reportlab. Run the Windows builder or install reportlab.") from exc
        styles = getSampleStyleSheet()
        story = []
        for paragraph in document.text.split("\n\n"):
            story.extend((Paragraph(escape(paragraph).replace("\n", "<br/>"), styles["BodyText"]), Spacer(1, 8)))
        SimpleDocTemplate(str(path), pagesize=letter, title=document.title, author="Dietrich AI Labs").build(story)

    def _schedule_autosave(self) -> None:
        if self._autosave_job:
            try:
                self.after_cancel(self._autosave_job)
            except Exception:
                pass
        self._autosave_job = self.after(max(10, int(self.settings.autosave_seconds)) * 1000, self._autosave)

    def _autosave(self) -> None:
        if self.dirty:
            document = self.serialize()
            self.recovery.save(
                RecoveryRecord(
                    recovery_id=self.recovery_id,
                    module="Writer",
                    title=document.title,
                    original_path=str(self.current_path or ""),
                    saved_at=datetime.now().isoformat(timespec="seconds"),
                    payload=document.to_dict(),
                )
            )
            self.status_var.set("Recovery copy saved")
        self._schedule_autosave()

    def open_find(self) -> None:
        if self.find_window and self.find_window.winfo_exists():
            self.find_window.lift()
            return
        win = tk.Toplevel(self)
        win.title("Find and Replace")
        win.transient(self.winfo_toplevel())
        win.resizable(False, False)
        self.find_window = win
        find_var, replace_var = tk.StringVar(), tk.StringVar()
        ttk.Label(win, text="Find").grid(row=0, column=0, sticky="w", padx=10, pady=(10, 4))
        find_entry = ttk.Entry(win, textvariable=find_var, width=38)
        find_entry.grid(row=0, column=1, padx=10, pady=(10, 4))
        ttk.Label(win, text="Replace").grid(row=1, column=0, sticky="w", padx=10, pady=4)
        ttk.Entry(win, textvariable=replace_var, width=38).grid(row=1, column=1, padx=10, pady=4)

        def find_next():
            self.text.tag_remove("search_hit", "1.0", "end")
            needle = find_var.get()
            if not needle:
                return
            start = self.text.index("insert")
            found = self.text.search(needle, start, stopindex="end", nocase=True) or self.text.search(needle, "1.0", stopindex=start, nocase=True)
            if found:
                end = f"{found}+{len(needle)}c"
                self.text.tag_configure("search_hit", background="#ffe28a")
                self.text.tag_add("search_hit", found, end)
                self.text.mark_set("insert", end)
                self.text.see(found)

        def replace_one():
            ranges = self.text.tag_ranges("search_hit")
            if ranges:
                self.text.delete(ranges[0], ranges[1])
                self.text.insert(ranges[0], replace_var.get())
                self.dirty = True
            find_next()

        def replace_all():
            needle = find_var.get()
            if not needle:
                return
            content = self.text.get("1.0", "end-1c")
            updated, count = re.subn(re.escape(needle), replace_var.get(), content, flags=re.IGNORECASE)
            if count:
                self.text.delete("1.0", "end")
                self.text.insert("1.0", updated)
                self.dirty = True
                messagebox.showinfo("Find and Replace", f"Replaced {count} occurrence(s).", parent=win)

        buttons = ttk.Frame(win)
        buttons.grid(row=2, column=0, columnspan=2, pady=10)
        ttk.Button(buttons, text="Find Next", command=find_next).pack(side="left", padx=3)
        ttk.Button(buttons, text="Replace", command=replace_one).pack(side="left", padx=3)
        ttk.Button(buttons, text="Replace All", command=replace_all).pack(side="left", padx=3)
        find_entry.focus_set()

    def recover_record(self, record: RecoveryRecord) -> None:
        if not self.confirm_discard_or_save():
            return
        self.load_document(LeanDocument.from_dict(record.payload), Path(record.original_path) if record.original_path else None)
        self.recovery_id = record.recovery_id
        self.dirty = True
        self._update_title_indicator()

    def print_document(self) -> None:
        if os.name != "nt":
            messagebox.showinfo("Print", "Printing is available in the Windows build.", parent=self)
            return
        temp = Path(tempfile.gettempdir()) / "LeanDesk_Print.rtf"
        write_text_document(self.serialize(), temp)
        try:
            os.startfile(str(temp), "print")
        except OSError as exc:
            messagebox.showerror("Print", str(exc), parent=self)


# LeanDesk correction-1: imported foreign sources are never ordinary-Save targets.
